from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCS_DIR = Path(
    r"C:\Users\gosfe\OneDrive\University\Third year\Second Semester\Machine Learning\Project Documnts"
)
OUTPUT_PATH = DOCS_DIR / "Retail Price Optimizer - Professor Questions Guide.docx"

TEAL = "1F6F78"
DARK = "17313A"
PALE = "E8F2F1"
GOLD = "F4E9CE"
LINE = "C7D5D6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = DARK) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_bidi(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(DARK if level == 1 else TEAL)


def add_question(doc: Document, number: int, question: str, english: str, arabic: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"{number}. {question}")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(DARK)

    english_p = doc.add_paragraph()
    english_p.paragraph_format.space_after = Pt(2)
    label = english_p.add_run("English: ")
    label.bold = True
    label.font.name = "Calibri"
    label.font.size = Pt(9.5)
    body = english_p.add_run(english)
    body.font.name = "Calibri"
    body.font.size = Pt(9.5)

    arabic_p = doc.add_paragraph()
    arabic_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    arabic_p.paragraph_format.space_after = Pt(5)
    set_bidi(arabic_p)
    ar_label = arabic_p.add_run("العربية: ")
    ar_label.bold = True
    ar_label.font.name = "Arial"
    ar_label.font.size = Pt(9.5)
    ar_body = arabic_p.add_run(arabic)
    ar_body.font.name = "Arial"
    ar_body.font.size = Pt(9.5)


def add_role_section(doc: Document, title: str, role_note: str, questions: list[tuple[str, str, str]]) -> None:
    add_heading(doc, title, level=2)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(6)
    run = note.add_run(role_note)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("5B636C")
    for idx, item in enumerate(questions, start=1):
        add_question(doc, idx, item[0], item[1], item[2])


general_questions = [
    (
        "What is your project about?",
        "Our project is a Retail Price Optimizer. It predicts product price in BRL using prepared e-commerce data, including product attributes, freight value, category, size, and demand-count features.",
        "مشروعنا هو نظام لتحسين وتسعير المنتجات في التجارة الإلكترونية. يتنبأ بسعر المنتج بالريال البرازيلي باستخدام بيانات مجهزة تشمل خصائص المنتج، قيمة الشحن، الفئة، الحجم، ومؤشرات الطلب.",
    ),
    (
        "What dataset did you use?",
        "We used the Brazilian E-Commerce Public Dataset by Olist. The final training file is dataset/feature_engineered_dataset.csv.",
        "استخدمنا مجموعة بيانات التجارة الإلكترونية البرازيلية من Olist. ملف التدريب النهائي هو dataset/feature_engineered_dataset.csv.",
    ),
    (
        "Why did you choose this dataset?",
        "It is realistic, structured, and suitable for regression because it contains product, freight, category, and price information from real e-commerce orders.",
        "اخترناها لأنها واقعية ومنظمة ومناسبة لمشكلة انحدار، لأنها تحتوي على معلومات المنتج والشحن والفئة والسعر من طلبات تجارة إلكترونية حقيقية.",
    ),
    (
        "What is the target variable?",
        "The target variable is price, which represents the item price in BRL.",
        "المتغير الهدف هو price، وهو يمثل سعر المنتج بالريال البرازيلي.",
    ),
    (
        "Is this classification or regression?",
        "It is regression because the model predicts a continuous numeric value: the product price.",
        "هذه مشكلة انحدار لأن النموذج يتنبأ بقيمة رقمية مستمرة وهي سعر المنتج.",
    ),
    (
        "What does optimization mean in your project?",
        "Optimization means selecting the model and parameter settings that reduce prediction error. We optimized using validation MAE through GridSearchCV, not direct profit maximization.",
        "التحسين في مشروعنا يعني اختيار النموذج والإعدادات التي تقلل خطأ التنبؤ. استخدمنا MAE في التحقق من خلال GridSearchCV، وليس تعظيم الربح بشكل مباشر.",
    ),
    (
        "Which models did you compare?",
        "We compared Linear Regression as a baseline, Random Forest Regressor, and Gradient Boosting Regressor.",
        "قارنّا بين Linear Regression كنموذج أساسي، و Random Forest Regressor، و Gradient Boosting Regressor.",
    ),
    (
        "Which model performed best?",
        "Random Forest Regressor performed best with MAE 22.54 BRL, RMSE 53.09 BRL, and R2 0.775 on the holdout test set.",
        "أفضل نموذج كان Random Forest Regressor وحقق MAE بقيمة 22.54، و RMSE بقيمة 53.09، و R2 بقيمة 0.775 على بيانات الاختبار.",
    ),
    (
        "Why did Random Forest perform well?",
        "Random Forest works well with tabular data and can capture nonlinear relationships between freight, product weight, category, dimensions, and price.",
        "أداء Random Forest كان جيدا لأنه مناسب للبيانات الجدولية ويستطيع التقاط العلاقات غير الخطية بين الشحن ووزن المنتج والفئة والأبعاد والسعر.",
    ),
    (
        "What do MAE, RMSE, and R2 mean?",
        "MAE is the average absolute error in BRL. RMSE penalizes larger errors more strongly. R2 shows how much price variation the model explains.",
        "MAE هو متوسط الخطأ المطلق بالريال البرازيلي. RMSE يعاقب الأخطاء الكبيرة بشكل أكبر. R2 يوضح مقدار التغير في السعر الذي يفسره النموذج.",
    ),
    (
        "How did you avoid data leakage?",
        "We removed target-derived columns such as total_order_item_value, freight_ratio, price_range, category average price, and category median price.",
        "تجنبنا تسرب البيانات بحذف الأعمدة المشتقة من السعر مثل total_order_item_value و freight_ratio و price_range ومتوسط ووسيط السعر للفئة.",
    ),
    (
        "What were the most important features?",
        "The strongest features included freight_value, product_weight_g, description length, product_demand_count, product_volume_cm3, and product dimensions.",
        "أهم الخصائص كانت freight_value و product_weight_g وطول الوصف وعدد الطلبات على المنتج وحجم المنتج وأبعاده.",
    ),
    (
        "What is the main limitation?",
        "The model predicts expected price from historical data. It does not include competitor prices, inventory, promotions, or real demand elasticity.",
        "النموذج يتنبأ بالسعر المتوقع من البيانات التاريخية، لكنه لا يحتوي على أسعار المنافسين أو المخزون أو العروض أو مرونة الطلب الحقيقية.",
    ),
    (
        "How can the project be improved?",
        "It can be improved by adding competitor prices, inventory, promotions, seasonality, product-level sales quantities, and a real profit optimization objective.",
        "يمكن تحسينه بإضافة أسعار المنافسين والمخزون والعروض والموسمية وكميات المبيعات لكل منتج وهدف تحسين للربح.",
    ),
]

role_questions = {
    "Abdulaziz Saud Aldossary - Leader / Model Training Lead": {
        "note": "Main responsibility: team leadership, GitHub/repo management, model-training workflow, methodology review, final validation, and submission preparation.",
        "questions": [
            (
                "What was your role as leader?",
                "I coordinated the team, maintained the GitHub repository, integrated the report sections, implemented and reviewed the main training workflow, validated results, and prepared the final submission package.",
                "دوري كقائد كان تنسيق عمل الفريق، إدارة مستودع GitHub، دمج أقسام التقرير، تنفيذ ومراجعة مسار تدريب النموذج، التحقق من النتائج، وتجهيز ملف التسليم النهائي.",
            ),
            (
                "What was your technical contribution?",
                "My technical contribution was leading the model-training workflow, reviewing the prepared data, training and comparing the models, checking leakage, and validating the final Random Forest result.",
                "مساهمتي التقنية كانت قيادة مسار تدريب النموذج، مراجعة البيانات المجهزة، تدريب ومقارنة النماذج، فحص تسرب البيانات، والتحقق من نتيجة Random Forest النهائية.",
            ),
            (
                "Why did you keep work on the model_training branch?",
                "We kept training artifacts on model_training so data-prep work stays separate and the final model outputs are easier to track and reproduce.",
                "استخدمنا فرع model_training حتى تبقى أعمال تجهيز البيانات منفصلة، وتكون مخرجات التدريب النهائي أسهل في التتبع وإعادة التشغيل.",
            ),
            (
                "How did you validate the final result?",
                "I checked holdout metrics, cross-validation results, feature-subset comparisons, feature importance, and a saved-model smoke test using the joblib pipeline.",
                "تحققت من النتيجة النهائية من خلال مقاييس الاختبار، نتائج التحقق المتقاطع، مقارنة مجموعات الخصائص، أهمية الخصائص، واختبار تشغيل النموذج المحفوظ بملف joblib.",
            ),
            (
                "What would you do differently next time?",
                "I would add stronger business data such as competitor prices, promotions, inventory, and demand elasticity so the model can move closer to true price optimization.",
                "في المرة القادمة سأضيف بيانات تجارية أقوى مثل أسعار المنافسين والعروض والمخزون ومرونة الطلب حتى يقترب النموذج من تحسين السعر الحقيقي.",
            ),
        ],
    },
    "Abdullah Almutairi - Data Preparation": {
        "note": "Main responsibility: dataset integration, cleaning, descriptive statistics, missing-value handling, and documentation of data-preparation decisions.",
        "questions": [
            (
                "What was your role in data preparation?",
                "My role was to help integrate the dataset, clean records, inspect descriptive statistics, handle missing values, and document preparation decisions.",
                "دوري كان المساعدة في دمج البيانات، تنظيف السجلات، فحص الإحصاءات الوصفية، معالجة القيم المفقودة، وتوثيق قرارات تجهيز البيانات.",
            ),
            (
                "Why is data cleaning important?",
                "Cleaning is important because missing, duplicated, or inconsistent records can make the model learn wrong patterns and reduce prediction quality.",
                "تنظيف البيانات مهم لأن القيم المفقودة أو المكررة أو غير المتسقة قد تجعل النموذج يتعلم أنماطا خاطئة وتقلل جودة التنبؤ.",
            ),
            (
                "What files represent the prepared data stages?",
                "The main files are cleaned_base_dataset.csv, price_outliers_removed_dataset.csv, and feature_engineered_dataset.csv.",
                "الملفات الأساسية لمراحل تجهيز البيانات هي cleaned_base_dataset.csv و price_outliers_removed_dataset.csv و feature_engineered_dataset.csv.",
            ),
            (
                "Why did the team remove outliers?",
                "Extreme price values can distort model training and make error metrics less representative, so outlier handling helped the model focus on realistic pricing patterns.",
                "القيم السعرية المتطرفة قد تؤثر على تدريب النموذج وتجعل المقاييس أقل تمثيلا، لذلك ساعدت معالجة القيم الشاذة النموذج على تعلم أنماط أسعار واقعية.",
            ),
            (
                "How did data preparation affect the final model?",
                "Good preparation made the training data more consistent, reduced noise, and gave the model cleaner features to learn from.",
                "تجهيز البيانات الجيد جعل بيانات التدريب أكثر اتساقا، وقلل الضوضاء، وقدم خصائص أوضح للنموذج ليتعلم منها.",
            ),
        ],
    },
    "Abdulaziz Alsmail - Literature Review / Feature-Engineering Support": {
        "note": "Main responsibility: literature review, gap identification, feature-engineering notes, final discussion support, and references.",
        "questions": [
            (
                "What was your role in the project?",
                "My role was supporting the literature review, identifying the project gap, documenting feature-engineering ideas, and helping connect results to the final discussion.",
                "دوري كان دعم مراجعة الأدبيات، تحديد فجوة المشروع، توثيق أفكار هندسة الخصائص، والمساعدة في ربط النتائج بالمناقشة النهائية.",
            ),
            (
                "What is the project gap?",
                "The gap is that simple fixed pricing rules do not fully use product, freight, category, and demand information. Our model uses these signals for data-driven price prediction.",
                "الفجوة هي أن قواعد التسعير الثابتة لا تستفيد بشكل كامل من معلومات المنتج والشحن والفئة والطلب. نموذجنا يستخدم هذه الإشارات للتنبؤ بالسعر بناء على البيانات.",
            ),
            (
                "Why is feature engineering useful?",
                "Feature engineering transforms raw data into more meaningful predictors, such as product volume, demand counts, freight level, and product size level.",
                "هندسة الخصائص مفيدة لأنها تحول البيانات الخام إلى مؤشرات أكثر معنى مثل حجم المنتج، عدد الطلبات، مستوى الشحن، ومستوى حجم المنتج.",
            ),
            (
                "How did references support the project?",
                "References supported the idea that machine learning can improve pricing decisions by learning nonlinear patterns from historical retail data.",
                "دعمت المراجع فكرة أن التعلم الآلي يمكن أن يحسن قرارات التسعير من خلال تعلم الأنماط غير الخطية من بيانات البيع التاريخية.",
            ),
            (
                "How do the results connect to the business discussion?",
                "The results show that freight, product scale, demand, and category signals influence predicted price, which supports more evidence-based pricing decisions.",
                "توضح النتائج أن الشحن وحجم المنتج والطلب والفئة تؤثر في السعر المتوقع، وهذا يدعم قرارات تسعير مبنية على الأدلة.",
            ),
        ],
    },
    "Mohammed Alkhamees - Model Testing / Optimization Support": {
        "note": "Main responsibility: model testing and optimization support, comparing model outputs, reviewing GridSearchCV parameter results, validating the saved pipeline, and documenting testing observations.",
        "questions": [
            (
                "What was your technical role?",
                "My role was model testing and optimization support. I helped compare model outputs, review GridSearchCV parameter results, validate the saved joblib pipeline, and test custom input predictions.",
                "دوري التقني كان اختبار النموذج ودعم التحسين. ساعدت في مقارنة مخرجات النماذج، مراجعة نتائج معاملات GridSearchCV، التحقق من ملف joblib، واختبار التنبؤ بقيم مخصصة.",
            ),
            (
                "What is GridSearchCV?",
                "GridSearchCV tests combinations of hyperparameters using cross-validation, then selects the setting with the best validation performance.",
                "GridSearchCV يختبر مجموعات مختلفة من معاملات النموذج باستخدام التحقق المتقاطع، ثم يختار الإعداد الذي يعطي أفضل أداء في التحقق.",
            ),
            (
                "What did the optimization support include?",
                "It included reviewing Random Forest and Gradient Boosting parameter options, checking selected settings, and comparing their final prediction errors.",
                "شمل دعم التحسين مراجعة اختيارات معاملات Random Forest و Gradient Boosting، فحص الإعدادات المختارة، ومقارنة أخطاء التنبؤ النهائية.",
            ),
            (
                "How did you test the saved model?",
                "I loaded models/best_price_model.joblib, checked the required 15 input columns, ran predictions on sample rows and manual inputs, and compared outputs with actual prices when available.",
                "اختبرت النموذج المحفوظ بتحميل models/best_price_model.joblib، وفحص الأعمدة المطلوبة وعددها 15، وتشغيل التنبؤ على عينات ومدخلات يدوية، ومقارنة النتائج بالأسعار الحقيقية عند توفرها.",
            ),
            (
                "Why is testing the saved model important?",
                "It proves the final model is reusable after training and that the preprocessing pipeline works correctly when new inputs are provided.",
                "اختبار النموذج المحفوظ مهم لأنه يثبت أن النموذج قابل لإعادة الاستخدام بعد التدريب وأن خطوات المعالجة المسبقة تعمل بشكل صحيح مع المدخلات الجديدة.",
            ),
        ],
    },
    "Hussain Alnasser - Evaluation / Visualization": {
        "note": "Main responsibility: evaluation tables, MAE/RMSE/R2 analysis, feature-importance visuals, and final presentation support.",
        "questions": [
            (
                "What was your role in evaluation?",
                "My role was preparing evaluation tables and figures, analyzing MAE, RMSE, and R2 results, creating feature-importance visuals, and supporting the final presentation.",
                "دوري كان إعداد جداول ورسومات التقييم، تحليل نتائج MAE و RMSE و R2، إنشاء رسوم أهمية الخصائص، ودعم العرض النهائي.",
            ),
            (
                "How do you explain MAE?",
                "MAE is the average absolute prediction error. For the final model, MAE 22.54 means the predicted price is off by about 22.54 BRL on average.",
                "MAE هو متوسط الخطأ المطلق في التنبؤ. في نموذجنا النهائي، قيمة 22.54 تعني أن التنبؤ يبتعد عن السعر الحقيقي بحوالي 22.54 ريال برازيلي في المتوسط.",
            ),
            (
                "How do you explain RMSE?",
                "RMSE gives more weight to large errors, so it helps identify whether the model makes big mistakes on unusual or expensive products.",
                "RMSE يعطي وزنا أكبر للأخطاء الكبيرة، لذلك يساعد في معرفة هل النموذج يرتكب أخطاء كبيرة في المنتجات غير المعتادة أو الغالية.",
            ),
            (
                "How do you explain R2?",
                "R2 shows how much variation in price is explained by the model. The final R2 of 0.775 means the model explains about 77.5% of test-set price variation.",
                "R2 يوضح مقدار التغير في السعر الذي يفسره النموذج. قيمة 0.775 تعني أن النموذج يفسر تقريبا 77.5% من تغير الأسعار في بيانات الاختبار.",
            ),
            (
                "Why are visualizations useful?",
                "Visualizations make the results easier to understand by showing model comparison, predicted versus actual prices, residual errors, and important features.",
                "الرسوم البيانية مفيدة لأنها تجعل النتائج أسهل للفهم من خلال عرض مقارنة النماذج، السعر المتوقع مقابل الحقيقي، الأخطاء المتبقية، وأهم الخصائص.",
            ),
        ],
    },
}


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Retail Price Optimizer\nProfessor Questions Guide")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("English and Arabic preparation answers for general project questions and role-specific questions")
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("5B636C")

    summary = doc.add_table(rows=1, cols=4)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(
        summary.rows[0].cells,
        ["Best model\nRandom Forest", "MAE\n22.54 BRL", "RMSE\n53.09 BRL", "R2\n0.775"],
        strict=True,
    ):
        set_cell_shading(cell, PALE)
        set_cell_text(cell, text, bold=True, color=TEAL)

    add_heading(doc, "Team Roles To Remember", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Student", "Role", "Main focus if professor asks"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, TEAL)
        set_cell_text(cell, text, bold=True, color="FFFFFF")
    rows = [
        ("Abdulaziz Saud Aldossary", "Leader / Model Training Lead", "Project coordination, GitHub, model-training workflow, final validation, submission package."),
        ("Abdullah Almutairi", "Data Preparation", "Dataset integration, cleaning, descriptive statistics, missing-value handling, preparation documentation."),
        ("Abdulaziz Alsmail", "Literature / Feature-Engineering Support", "Literature review, gap identification, feature-engineering notes, discussion and references."),
        ("Mohammed Alkhamees", "Model Testing / Optimization Support", "Test trained models, review GridSearchCV results, validate saved pipeline, custom prediction tests."),
        ("Hussain Alnasser", "Evaluation / Visualization", "Evaluation tables, MAE/RMSE/R2 analysis, feature-importance visuals, presentation support."),
    ]
    for row_idx, row_data in enumerate(rows, start=1):
        cells = table.add_row().cells
        for col_idx, text in enumerate(row_data):
            set_cell_text(cells[col_idx], text, color=DARK)
            if row_idx % 2 == 1:
                set_cell_shading(cells[col_idx], "F3F8F8")

    add_heading(doc, "General Project Questions", level=1)
    for idx, item in enumerate(general_questions, start=1):
        add_question(doc, idx, item[0], item[1], item[2])

    doc.add_page_break()
    add_heading(doc, "Role-Specific Questions", level=1)
    for title, info in role_questions.items():
        add_role_section(doc, title, info["note"], info["questions"])

    add_heading(doc, "Safe Answer Structure", level=1)
    add_question(
        doc,
        1,
        "What should I do if I get a question outside my exact role?",
        "Start with your role, connect it to the project pipeline, then mention the teammate or artifact that owns the detailed part. Example: My role was model testing, so I can explain how we loaded and validated the saved model. The main training notebook contains the deeper training details.",
        "ابدأ بذكر دورك، ثم اربطه بمسار المشروع، ثم اذكر العضو أو الملف المسؤول عن التفاصيل. مثال: دوري كان اختبار النموذج، لذلك أستطيع شرح كيفية تحميل النموذج المحفوظ والتحقق منه، أما تفاصيل التدريب العميقة فهي موجودة في دفتر التدريب الرئيسي.",
    )
    add_question(
        doc,
        2,
        "What is the shortest full-project answer?",
        "We used the Olist e-commerce dataset to predict product price using regression. After cleaning and feature engineering, Random Forest gave the best result with MAE 22.54 BRL and R2 0.775.",
        "استخدمنا بيانات Olist للتجارة الإلكترونية للتنبؤ بسعر المنتج باستخدام الانحدار. بعد تنظيف البيانات وهندسة الخصائص، حقق Random Forest أفضل نتيجة بقيمة MAE 22.54 و R2 0.775.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Retail Price Optimizer | ARTI 308 Machine Learning")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("6B747B")

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
