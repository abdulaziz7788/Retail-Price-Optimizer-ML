# -*- coding: utf-8 -*-
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(
    r"C:\Users\gosfe\OneDrive\University\Third year\Second Semester\Machine Learning\Project repo"
)
DOCS_DIR = Path(
    r"C:\Users\gosfe\OneDrive\University\Third year\Second Semester\Machine Learning\Project Documnts"
)
OUTPUT_PATH = DOCS_DIR / "Retail Price Optimizer - Professor Questions Guide.docx"

TEAL = "1F6F78"
DARK = "17313A"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
PALE_TEAL = "E8F2F1"
PALE_GOLD = "F4E9CE"
PALE_GRAY = "F4F6F9"
LINE = "C7D5D6"
MUTED = "5B636C"
WHITE = "FFFFFF"


@dataclass(frozen=True)
class Question:
    question: str
    english: str
    arabic: str
    code: str


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def dxa(inches: float) -> int:
    return int(round(inches * 1440))


def set_shading(element, fill: str) -> None:
    p_pr = element.get_or_add_pPr() if element.tag.endswith("}p") else element.get_or_add_tcPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_bidi(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_run(paragraph, text: str, *, bold=False, italic=False, size=10.5, color=DARK, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    return run


def add_plain_paragraph(doc: Document, text: str = "", *, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if text:
        add_run(p, text)
    return p


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_run(
        p,
        "Retail Price Optimizer - Professor Questions Guide",
        bold=True,
        size=20,
        color=DARK,
    )
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    add_run(
        p2,
        "Technical edition with English/Arabic answers, code evidence, and likely discussion questions",
        size=11,
        color=MUTED,
    )


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.color.rgb = rgb(DARK if level == 1 else TEAL)
        run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 12)


def add_label_text(doc: Document, label: str, text: str, *, label_color=TEAL) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_run(p, label, bold=True, size=9.7, color=label_color)
    add_run(p, text, size=9.7, color=DARK)


def add_arabic_answer(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_bidi(p)
    add_run(p, "العربية: ", bold=True, size=9.7, color=TEAL, font="Arial")
    add_run(p, text, size=9.7, color=DARK, font="Arial")


def add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph(style="CodeBlock")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    set_shading(p._p, PALE_GRAY)
    add_run(p, "Code / technical evidence:\n", bold=True, size=8.3, color=TEAL, font="Consolas")
    add_run(p, code.strip(), size=8.1, color=DARK, font="Consolas")


def add_question(doc: Document, number: int, item: Question) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    add_run(p, f"Q{number}. {item.question}", bold=True, size=11.2, color=DARK)
    add_label_text(doc, "English: ", item.english)
    add_arabic_answer(doc, item.arabic)
    add_code_block(doc, item.code)


def add_key_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [dxa(1.72), dxa(4.78)]
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Value / Evidence"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(9.2)
                run.font.color.rgb = rgb(DARK)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for idx, cell in enumerate(cells):
            set_cell_margins(cell)
            if idx == 0:
                set_cell_shading(cell, PALE_TEAL)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.0)
                    run.font.color.rgb = rgb(DARK)
                    if idx == 0:
                        run.bold = True
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, DARK, 18, 10),
        ("Heading 2", 13, TEAL, 14, 7),
        ("Heading 3", 12, BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "CodeBlock" not in doc.styles:
        code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["CodeBlock"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.font.size = Pt(8.1)
    code_style.font.color.rgb = rgb(DARK)
    code_style.paragraph_format.line_spacing = 1.05
    code_style.paragraph_format.space_after = Pt(8)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "Retail Price Optimizer | ARTI 308 | Professor Questions Technical Guide", size=8.5, color=MUTED)


general_questions = [
    Question(
        "What is the project about?",
        "It is a Retail Price Optimizer for Brazilian e-commerce. The model predicts item price in BRL from product, freight, category, size, and demand-related features.",
        "المشروع هو نظام لتحسين وتسعير منتجات التجارة الإلكترونية البرازيلية. النموذج يتنبأ بسعر المنتج بالريال البرازيلي اعتمادا على خصائص المنتج والشحن والفئة والحجم ومؤشرات الطلب.",
        'Project file: notebooks/03_model_training_results.ipynb\nTarget: price\nBest model: Random Forest Regressor\nFinal metric: R2 = 0.775, MAE = 22.54 BRL',
    ),
    Question(
        "What dataset did you use?",
        "We used the Olist Brazilian E-Commerce Public Dataset. The modeling file used by the training branch is dataset/feature_engineered_dataset.csv.",
        "استخدمنا مجموعة بيانات Olist للتجارة الإلكترونية البرازيلية. ملف التدريب المستخدم في فرع التدريب هو dataset/feature_engineered_dataset.csv.",
        'import pandas as pd\n\ndf = pd.read_csv("dataset/feature_engineered_dataset.csv")\nprint(df.shape)  # 107,529 rows after preparation',
    ),
    Question(
        "Why did you choose this dataset?",
        "It is realistic, structured, and suitable for regression because it contains real transaction prices with product attributes, freight cost, category information, and order behavior.",
        "اخترناها لأنها واقعية ومنظمة ومناسبة للانحدار، لأنها تحتوي على أسعار معاملات حقيقية مع خصائص المنتج وتكلفة الشحن وفئة المنتج وسلوك الطلبات.",
        'Useful signals in the prepared file:\nprice, freight_value, product_weight_g, product_volume_cm3,\nproduct_category_name_english, product_demand_count',
    ),
    Question(
        "Is this a classification or regression problem?",
        "It is a regression problem because the model predicts a continuous numeric value: product price.",
        "هذه مشكلة انحدار لأن النموذج يتنبأ بقيمة رقمية مستمرة وهي سعر المنتج.",
        'X = df.drop(columns=["price"])\ny = df["price"]\n# y is continuous, so we use regressors, not classifiers.',
    ),
    Question(
        "What is the target variable?",
        "The target variable is price, representing item price in BRL.",
        "المتغير الهدف هو price ويمثل سعر المنتج بالريال البرازيلي.",
        'target = "price"\ny = df[target]\nprint(y.mean(), y.median(), y.max())\n# Mean 107.02, median 72.99, max 887.00 BRL',
    ),
    Question(
        "What are the main input features?",
        "The final model uses 15 cleaned input columns, including freight value, product weight, dimensions, category, name length, description length, photo count, demand count, and prepared categorical levels.",
        "النموذج النهائي يستخدم 15 خاصية نظيفة، تشمل قيمة الشحن ووزن المنتج والأبعاد والفئة وطول الاسم والوصف وعدد الصور وعدد الطلبات ومستويات فئوية مجهزة.",
        'final_features = [c for c in df.columns if c != "price"]\nprint(len(final_features))  # 15 input columns',
    ),
    Question(
        "What does optimization mean in your project?",
        "In this project, optimization means selecting model settings that reduce prediction error. It is model optimization using validation MAE, not direct profit maximization.",
        "في هذا المشروع التحسين يعني اختيار إعدادات النموذج التي تقلل خطأ التنبؤ. هو تحسين للنموذج باستخدام MAE في التحقق، وليس تعظيم الربح بشكل مباشر.",
        'from sklearn.model_selection import GridSearchCV\n\ngrid = GridSearchCV(\n    estimator=rf_pipeline,\n    param_grid=rf_param_grid,\n    scoring="neg_mean_absolute_error",\n    cv=3,\n)',
    ),
    Question(
        "What is the difference between price prediction and true price optimization?",
        "Price prediction estimates a likely price based on historical data. True price optimization would also need demand elasticity, profit margin, inventory, competitor prices, and business constraints.",
        "التنبؤ بالسعر يقدر السعر المتوقع من البيانات التاريخية. أما تحسين السعر الحقيقي فيحتاج مرونة الطلب وهامش الربح والمخزون وأسعار المنافسين وقيود العمل.",
        'Current objective: minimize prediction error\nFuture objective: maximize profit = (price - cost) * expected_demand(price)\n# The current dataset does not contain cost or elasticity fields.',
    ),
    Question(
        "How did you split the data?",
        "We used a holdout train-test split. The report uses 75,270 training rows and 32,259 test rows.",
        "استخدمنا تقسيم تدريب واختبار. التقرير يستخدم 75,270 صفا للتدريب و32,259 صفا للاختبار.",
        'from sklearn.model_selection import train_test_split\n\nX_train, X_test, y_train, y_test = train_test_split(\n    X, y, test_size=0.30, random_state=42\n)',
    ),
    Question(
        "How did you avoid data leakage?",
        "We removed columns derived from price before training, such as total order value, freight ratio, price range, and category price aggregates.",
        "تجنبنا تسرب البيانات بحذف الأعمدة المشتقة من السعر قبل التدريب، مثل إجمالي قيمة الطلب ونسبة الشحن ونطاق السعر ومتوسطات سعر الفئة.",
        'leakage_cols = [\n    "total_order_item_value", "freight_ratio", "price_range",\n    "category_avg_price", "category_median_price"\n]\nX = df.drop(columns=["price"] + [c for c in leakage_cols if c in df.columns])',
    ),
    Question(
        "Why is target leakage dangerous?",
        "Leakage makes the model look much better than it really is because it gives the model information that would not be available for a new product prediction.",
        "تسرب الهدف خطير لأنه يجعل النموذج يبدو أفضل من الواقع، لأنه يعطيه معلومات لن تكون متاحة عند التنبؤ بسعر منتج جديد.",
        'Bad example:\nX_bad = df[["category_avg_price", "freight_ratio"]]\n# These can indirectly reveal price, so they must not be used for fair testing.',
    ),
    Question(
        "How did you handle missing values?",
        "The pipeline uses imputers: numeric columns use median values, and categorical columns use the most frequent category before encoding.",
        "المسار يستخدم تعويض القيم المفقودة: الأعمدة الرقمية تستخدم الوسيط، والأعمدة الفئوية تستخدم القيمة الأكثر تكرارا قبل الترميز.",
        'from sklearn.impute import SimpleImputer\n\nnumeric_pipe = SimpleImputer(strategy="median")\ncategorical_pipe = SimpleImputer(strategy="most_frequent")',
    ),
    Question(
        "How did you handle categorical features?",
        "Categorical columns are converted using OneHotEncoder, with unknown categories ignored so the saved model can still run on new values.",
        "الأعمدة الفئوية تم تحويلها باستخدام OneHotEncoder مع تجاهل الفئات غير المعروفة حتى يعمل النموذج المحفوظ مع قيم جديدة.",
        'from sklearn.preprocessing import OneHotEncoder\n\nencoder = OneHotEncoder(handle_unknown="ignore", sparse_output=False)',
    ),
    Question(
        "Why did you use a Pipeline?",
        "A Pipeline keeps preprocessing and the model together, so training and prediction use the same steps. This also makes the saved joblib file easier to reuse.",
        "استخدمنا Pipeline لأنه يجمع تجهيز البيانات والنموذج في مسار واحد، لذلك يستخدم التدريب والتنبؤ نفس الخطوات. وهذا يسهل إعادة استخدام ملف joblib.",
        'from sklearn.pipeline import Pipeline\n\nmodel_pipe = Pipeline([\n    ("preprocess", preprocess),\n    ("model", RandomForestRegressor(random_state=42))\n])',
    ),
    Question(
        "Which models did you compare?",
        "We compared Linear Regression as a baseline, Random Forest Regressor, and Gradient Boosting Regressor.",
        "قارنّا بين Linear Regression كنموذج أساسي، وRandom Forest Regressor، وGradient Boosting Regressor.",
        'models = {\n    "Linear Regression Baseline": linear_pipeline,\n    "Random Forest Regressor": rf_pipeline,\n    "Gradient Boosting Regressor": gb_pipeline,\n}',
    ),
    Question(
        "Why include Linear Regression if it performs worse?",
        "Linear Regression gives a simple baseline. It helps show that the nonlinear ensemble model truly adds value over a basic model.",
        "أضفنا Linear Regression كنقطة أساس بسيطة. هذا يوضح أن النموذج التجميعي غير الخطي يضيف قيمة فعلية مقارنة بنموذج بسيط.",
        'from sklearn.linear_model import LinearRegression\n\nbaseline = Pipeline([\n    ("preprocess", preprocess),\n    ("model", LinearRegression())\n])',
    ),
    Question(
        "Why use Random Forest?",
        "Random Forest works well on tabular data, captures nonlinear relationships, handles mixed features after preprocessing, and reduces variance by averaging many trees.",
        "استخدمنا Random Forest لأنه مناسب للبيانات الجدولية، يلتقط العلاقات غير الخطية، يتعامل مع الخصائص المختلطة بعد التجهيز، ويقلل التباين عبر متوسط عدة أشجار.",
        'from sklearn.ensemble import RandomForestRegressor\n\nrf = RandomForestRegressor(\n    n_estimators=140,\n    max_depth=None,\n    random_state=42,\n    n_jobs=-1,\n)',
    ),
    Question(
        "Why use Gradient Boosting?",
        "Gradient Boosting was tested because it learns sequentially from previous residual errors and is a common strong model for structured regression.",
        "اختبرنا Gradient Boosting لأنه يتعلم بشكل متسلسل من أخطاء النماذج السابقة، وهو نموذج قوي شائع في الانحدار على البيانات المنظمة.",
        'from sklearn.ensemble import GradientBoostingRegressor\n\ngb = GradientBoostingRegressor(\n    n_estimators=140,\n    learning_rate=0.10,\n    random_state=42,\n)',
    ),
    Question(
        "Which model performed best?",
        "Random Forest performed best on the holdout set with MAE 22.54 BRL, RMSE 53.09 BRL, and R2 0.775.",
        "أفضل نموذج كان Random Forest على بيانات الاختبار، حيث حقق MAE بقيمة 22.54 وRMSE بقيمة 53.09 وR2 بقيمة 0.775.",
        'pd.read_csv("reports/model_performance.csv")\n# Random Forest: MAE 22.54, RMSE 53.09, R2 0.775',
    ),
    Question(
        "How accurate is the model?",
        "The most direct answer is that the average absolute error is about 22.54 BRL, and the model explains about 77.5 percent of price variance on the holdout set.",
        "الإجابة المباشرة أن متوسط الخطأ المطلق حوالي 22.54 ريال برازيلي، والنموذج يفسر حوالي 77.5% من تباين السعر على بيانات الاختبار.",
        'MAE = 22.54396893058473\nR2 = 0.7753093018276964\nprint(f"Average error: {MAE:.2f} BRL, explained variance: {R2:.3f}")',
    ),
    Question(
        "What does MAE mean?",
        "MAE is the average absolute difference between actual and predicted price. It is easy to explain because it is measured in BRL.",
        "MAE هو متوسط الفرق المطلق بين السعر الحقيقي والسعر المتوقع. يسهل شرحه لأنه مقاس بالريال البرازيلي.",
        'from sklearn.metrics import mean_absolute_error\n\nmae = mean_absolute_error(y_test, y_pred)',
    ),
    Question(
        "What does RMSE mean?",
        "RMSE is another error metric that punishes large mistakes more strongly than MAE.",
        "RMSE هو مقياس خطأ آخر يعاقب الأخطاء الكبيرة أكثر من MAE.",
        'from sklearn.metrics import mean_squared_error\n\nrmse = mean_squared_error(y_test, y_pred, squared=False)',
    ),
    Question(
        "What does R2 mean?",
        "R2 shows the proportion of price variance explained by the model. An R2 of 0.775 means the model captures a strong share of the pricing pattern.",
        "R2 يوضح نسبة تباين السعر التي يفسرها النموذج. قيمة 0.775 تعني أن النموذج يلتقط جزءا قويا من نمط التسعير.",
        'from sklearn.metrics import r2_score\n\nr2 = r2_score(y_test, y_pred)',
    ),
    Question(
        "Why was Random Forest much better than Gradient Boosting here?",
        "On this feature set and tuning grid, Random Forest handled the mixed feature space better. Gradient Boosting may need deeper tuning, different loss functions, or more feature transformation.",
        "في هذه الخصائص وإعدادات البحث، تعامل Random Forest مع مساحة الخصائص المختلطة بشكل أفضل. قد يحتاج Gradient Boosting إلى ضبط أعمق أو دوال خسارة مختلفة أو تحويلات خصائص أكثر.",
        'Holdout results:\nRandom Forest MAE = 22.54\nGradient Boosting MAE = 52.89\nLinear Regression MAE = 57.86',
    ),
    Question(
        "How did you use cross-validation?",
        "We used cross-validation to check whether the model result is stable and not only good on one split. Random Forest had better CV MAE than Gradient Boosting.",
        "استخدمنا التحقق المتقاطع للتأكد من أن النتيجة مستقرة وليست جيدة فقط على تقسيم واحد. Random Forest حقق MAE أفضل في التحقق المتقاطع من Gradient Boosting.",
        'pd.read_csv("reports/cross_validation_results.csv")\n# Random Forest CV MAE mean: 39.29\n# Gradient Boosting CV MAE mean: 54.04',
    ),
    Question(
        "What were the most important features?",
        "The strongest feature signals were freight_value, product_weight_g, product description length, product_demand_count, product_volume_cm3, and product dimensions.",
        "أقوى الخصائص كانت freight_value وproduct_weight_g وطول وصف المنتج وproduct_demand_count وproduct_volume_cm3 وأبعاد المنتج.",
        'imp = pd.read_csv("reports/feature_importance.csv")\nprint(imp.head(8))',
    ),
    Question(
        "Why is freight_value important?",
        "Freight cost reflects logistics burden and often correlates with product size, weight, and operational cost, which can influence final item price.",
        "قيمة الشحن مهمة لأنها تعكس عبء التوصيل وغالبا ترتبط بحجم المنتج ووزنه والتكلفة التشغيلية، وهذا قد يؤثر على السعر النهائي.",
        'df[["price", "freight_value"]].corr()\n# Correlation with price was about 0.38',
    ),
    Question(
        "Why is product_weight_g important?",
        "Heavier products often cost more to produce, handle, and ship. In this dataset, product weight had a strong positive relationship with price.",
        "وزن المنتج مهم لأن المنتجات الأثقل غالبا تكون أعلى تكلفة في التصنيع والمناولة والشحن. في هذه البيانات كان للوزن علاقة موجبة واضحة مع السعر.",
        'df[["price", "product_weight_g"]].corr()\n# Correlation with price was about 0.37',
    ),
    Question(
        "Why did all 15 features give the best result?",
        "Keeping all cleaned features preserved category and engineered information. Removing category encodings or using only numeric features reduced model performance.",
        "استخدام كل الخصائص النظيفة حافظ على معلومات الفئة والخصائص المهندسة. حذف ترميزات الفئة أو استخدام الخصائص الرقمية فقط خفض أداء النموذج.",
        'pd.read_csv("reports/feature_subset_results.csv")\n# All 15 features: R2 0.775\n# Top numeric only: R2 0.711',
    ),
    Question(
        "Which engineered features came from the data-prep work?",
        "The model uses the prepared engineered file from the data-prep branch. Examples include product volume, demand count, freight level, and product size level.",
        "النموذج استخدم ملف الخصائص المهندسة الجاهز من فرع تجهيز البيانات. أمثلة ذلك حجم المنتج وعدد الطلبات ومستوى الشحن ومستوى حجم المنتج.",
        'Prepared columns include:\nproduct_volume_cm3\nproduct_demand_count\nfreight_level\nproduct_size_level',
    ),
    Question(
        "Did model training add new features?",
        "No. The model-training branch uses the feature-engineered dataset prepared by the team and removes only leakage columns before training.",
        "لا. فرع تدريب النموذج يستخدم ملف الخصائص المهندسة الذي جهزه الفريق، ويحذف فقط الأعمدة التي تسبب تسرب الهدف قبل التدريب.",
        'model_training uses:\ndataset/feature_engineered_dataset.csv\nLeakage columns are excluded before fitting the model.',
    ),
    Question(
        "How can the saved model be tested?",
        "The saved joblib pipeline can be loaded and used directly. The repository also has scripts/run_saved_model.py for a quick smoke test.",
        "يمكن اختبار نموذج joblib المحفوظ بتحميله واستخدامه مباشرة. يوجد أيضا scripts/run_saved_model.py لاختبار سريع.",
        'python scripts/run_saved_model.py --rows 3\n\n# Or in Python:\nmodel = joblib.load("models/best_price_model.joblib")\nmodel.predict(sample_rows)',
    ),
    Question(
        "How can a team member test custom values?",
        "Create a DataFrame or CSV with the same input columns used during training, then call model.predict on it.",
        "يمكن إنشاء DataFrame أو ملف CSV بنفس أعمدة الإدخال المستخدمة في التدريب، ثم تشغيل model.predict عليه.",
        'import joblib, pandas as pd\n\nmodel = joblib.load("models/best_price_model.joblib")\nnew_rows = pd.read_csv("dataset/manual_test_input.csv")\npreds = model.predict(new_rows)',
    ),
    Question(
        "What happens if a new category appears during prediction?",
        "The encoder uses handle_unknown='ignore', so unseen categories do not crash prediction. They are encoded as all-zero category indicators.",
        "إذا ظهرت فئة جديدة في التنبؤ فلن يتوقف النموذج، لأن الترميز يستخدم handle_unknown='ignore'. سيتم تمثيل الفئة الجديدة بأصفار في مؤشرات الفئات.",
        'OneHotEncoder(handle_unknown="ignore")\n# This is important for testing products with categories not seen during fitting.',
    ),
    Question(
        "What are the main project limitations?",
        "The data is observational. It does not include competitor prices, inventory, promotions, profit margins, or controlled repeated pricing experiments.",
        "أهم القيود أن البيانات تاريخية رصدية. لا تحتوي على أسعار المنافسين أو المخزون أو العروض أو هوامش الربح أو تجارب تسعير متكررة ومضبوطة.",
        'Missing business variables:\ncompetitor_price, inventory_level, promotion_flag,\nunit_cost, profit_margin, demand_elasticity',
    ),
    Question(
        "How could the project be improved?",
        "Future work should add competitor prices, inventory, promotions, seasonality, cost/margin data, and a demand model to move closer to real price optimization.",
        "يمكن تحسين المشروع بإضافة أسعار المنافسين والمخزون والعروض والموسمية وبيانات التكلفة والهامش ونموذج للطلب حتى نقترب من تحسين السعر الحقيقي.",
        'Future objective example:\nexpected_profit = (candidate_price - unit_cost) * demand_model.predict(candidate_price)\nbest_price = candidate_prices[expected_profit.argmax()]',
    ),
    Question(
        "How do you prove the project is reproducible?",
        "The repository includes datasets, notebooks, scripts, saved model, metrics CSV files, HTML results, and a README explaining how to navigate the branch.",
        "إثبات قابلية إعادة التشغيل يكون بوجود البيانات والدفاتر والسكريبتات والنموذج المحفوظ وملفات المقاييس ونتائج HTML وملف README يشرح تنظيم الفرع.",
        'Repository evidence:\nREADME.md\nnotebooks/03_model_training_results.ipynb\nscripts/run_saved_model.py\nreports/model_performance.csv\nmodels/best_price_model.joblib',
    ),
]


role_questions = {
    "Abdulaziz Saud Aldossary - Leader / Model Training Workflow": [
        Question(
            "What was your role as leader?",
            "I coordinated the team, maintained the repository structure, integrated final materials, led the model-training workflow, checked validation, and prepared submission artifacts.",
            "دوري كقائد كان تنسيق الفريق، إدارة هيكل المستودع، دمج مواد التسليم، قيادة مسار تدريب النموذج، التحقق من النتائج، وتجهيز ملفات التسليم.",
            'Leader evidence:\nREADME.md\nnotebooks/03_model_training_results.ipynb\nreports/model_training_metrics.json\nProject Documnts/ARTI 308 Final Project Report - Retail Price Optimizer.docx',
        ),
        Question(
            "What did you do technically in model training?",
            "I used the prepared feature-engineered data, removed leakage columns, trained baseline and ensemble regressors, compared metrics, saved the best pipeline, and validated it with a smoke test.",
            "تقنيا استخدمت ملف الخصائص المهندسة، حذفت أعمدة التسرب، دربت نموذج أساس ونماذج تجميعية، قارنت المقاييس، حفظت أفضل Pipeline، واختبرته اختبارا سريعا.",
            'best_model.fit(X_train, y_train)\ny_pred = best_model.predict(X_test)\njoblib.dump(best_model, "models/best_price_model.joblib")',
        ),
        Question(
            "Why keep training work on the model_training branch?",
            "The branch separates model-training artifacts from data-preparation work, so the data-prep branch remains aligned with the member who prepared the dataset.",
            "استخدمنا فرع model_training لفصل ملفات التدريب عن عمل تجهيز البيانات، حتى يبقى فرع data-prep مرتبطا بعضو الفريق الذي جهز البيانات.",
            'git branch\n# model_training\n# Training artifacts stay here; prepared datasets are consumed, not rewritten.',
        ),
        Question(
            "What would you say if asked whether this is a real optimizer?",
            "I would say it is a prediction-based price recommendation model. It optimizes model settings for lower prediction error, but true profit optimization requires more business variables.",
            "سأقول إنه نموذج توصية سعرية قائم على التنبؤ. هو يحسن إعدادات النموذج لتقليل خطأ التنبؤ، لكن تحسين الربح الحقيقي يحتاج متغيرات تجارية إضافية.",
            'Current: GridSearchCV(scoring="neg_mean_absolute_error")\nFuture: optimize profit under demand, cost, inventory, and competitor constraints.',
        ),
    ],
    "Abdullah Almutairi - Data Preparation": [
        Question(
            "What was your technical role?",
            "My role was dataset integration, cleaning, descriptive statistics, missing-value handling, and documenting why each data-preparation decision was made.",
            "دوري التقني كان دمج البيانات وتنظيفها، الإحصاءات الوصفية، معالجة القيم المفقودة، وتوثيق سبب كل قرار في تجهيز البيانات.",
            'Data-prep evidence files:\ndataset/cleaned_base_dataset.csv\ndataset/price_outliers_removed_dataset.csv\ndataset/feature_engineered_dataset.csv',
        ),
        Question(
            "How do you check missing values?",
            "We inspect null counts before modeling and then the training pipeline handles remaining missing values using imputers.",
            "نفحص عدد القيم المفقودة قبل النمذجة، ثم يتعامل مسار التدريب مع القيم المتبقية باستخدام التعويض.",
            'df = pd.read_csv("dataset/feature_engineered_dataset.csv")\nmissing = df.isna().sum().sort_values(ascending=False)\nprint(missing.head(10))',
        ),
        Question(
            "Why remove or handle price outliers?",
            "Extreme prices can distort training and inflate error metrics. Handling outliers helps the model learn more representative retail pricing patterns.",
            "الأسعار المتطرفة قد تشوه التدريب وتزيد مقاييس الخطأ. التعامل معها يساعد النموذج على تعلم أنماط تسعير أكثر تمثيلا.",
            'df["price"].describe()\n# Final report uses max price 887.00 BRL after preparation/outlier handling.',
        ),
    ],
    "Abdulaziz Alsmail - Literature Review / Feature Engineering Notes": [
        Question(
            "What was the academic gap?",
            "The gap is that simple pricing rules and moving averages often ignore nonlinear relationships between logistics, product scale, category, and demand signals.",
            "الفجوة الأكاديمية أن قواعد التسعير البسيطة والمتوسطات المتحركة غالبا تتجاهل العلاقات غير الخطية بين اللوجستيات وحجم المنتج والفئة وإشارات الطلب.",
            'Report evidence:\nSection: Literature Review and Project Gap\nTechnical link: tree ensembles capture nonlinear feature interactions.',
        ),
        Question(
            "Why is feature engineering important?",
            "Feature engineering turns raw columns into more meaningful predictors, such as product volume, demand count, freight level, and size level.",
            "هندسة الخصائص مهمة لأنها تحول الأعمدة الخام إلى مؤشرات أكثر معنى مثل حجم المنتج وعدد الطلبات ومستوى الشحن ومستوى الحجم.",
            'df["product_volume_cm3"] = (\n    df["product_length_cm"] * df["product_height_cm"] * df["product_width_cm"]\n)\n# Also prepared: product_demand_count, freight_level, product_size_level',
        ),
        Question(
            "How do the features connect to the literature?",
            "The features represent operational pricing factors: logistics cost, product scale, category identity, and demand volume. These address the gap in fixed-rule pricing.",
            "الخصائص تمثل عوامل تشغيلية للتسعير: تكلفة اللوجستيات، حجم المنتج، هوية الفئة، وحجم الطلب. هذه تعالج ضعف التسعير بالقواعد الثابتة.",
            'Feature groups:\nlogistics: freight_value\nscale: product_weight_g, product_volume_cm3\ndemand: product_demand_count\ncategory: product_category_name_english',
        ),
    ],
    "Mohammed Alkhamees - Model Testing / Optimization Support": [
        Question(
            "What was your technical role?",
            "My role was supporting model testing and optimization: checking GridSearchCV settings, reviewing parameter choices, validating saved-pipeline predictions, and testing custom inputs.",
            "دوري التقني كان دعم اختبار النموذج وتحسينه: مراجعة إعدادات GridSearchCV، مراجعة اختيار المعاملات، التحقق من تنبؤات Pipeline المحفوظ، واختبار مدخلات مخصصة.",
            'Optimization/test evidence:\nreports/model_performance.csv\nreports/cross_validation_results.csv\nscripts/run_saved_model.py\ndataset/manual_test_input.csv',
        ),
        Question(
            "What is GridSearchCV?",
            "GridSearchCV tries combinations of hyperparameters and uses cross-validation to choose the combination with the best validation score.",
            "GridSearchCV يجرب مجموعات مختلفة من المعاملات ويستخدم التحقق المتقاطع لاختيار المجموعة صاحبة أفضل نتيجة تحقق.",
            'param_grid = {\n    "model__n_estimators": [100, 140],\n    "model__max_depth": [None, 15, 25]\n}\ngrid = GridSearchCV(rf_pipeline, param_grid, scoring="neg_mean_absolute_error", cv=3)',
        ),
        Question(
            "How can you prove the saved model works?",
            "Load the joblib file and predict using a small valid input sample. The script run_saved_model.py performs this test from the command line.",
            "يمكن إثبات عمل النموذج بتحميل ملف joblib والتنبؤ بعينة إدخال صحيحة. سكريبت run_saved_model.py ينفذ هذا الاختبار من سطر الأوامر.",
            'python scripts/run_saved_model.py --rows 3\n\n# Expected behavior: prints 3 predicted prices in BRL without errors.',
        ),
        Question(
            "How would you test your own values?",
            "Use the same feature columns as training, place the values in a one-row DataFrame or CSV, then call model.predict.",
            "لاختبار قيم جديدة نستخدم نفس أعمدة التدريب ونضع القيم في DataFrame أو CSV من صف واحد ثم نشغل model.predict.",
            'sample = pd.DataFrame([{\n    "freight_value": 25.0,\n    "product_weight_g": 800,\n    "product_volume_cm3": 12000,\n    "product_category_name_english": "computers_accessories"\n}])\nmodel.predict(sample)',
        ),
    ],
    "Hussain Alnasser - Evaluation / Visualization": [
        Question(
            "What was your technical role?",
            "My role was evaluation design: preparing metric tables, analyzing MAE/RMSE/R2, reviewing cross-validation, and creating feature-importance/result visuals.",
            "دوري التقني كان تصميم التقييم: إعداد جداول المقاييس، تحليل MAE وRMSE وR2، مراجعة التحقق المتقاطع، وإنشاء رسوم أهمية الخصائص والنتائج.",
            'Evaluation evidence:\nreports/model_performance.csv\nreports/feature_importance.csv\nreports/figures/feature_importance.png\nreports/figures/model_error_comparison.png',
        ),
        Question(
            "How did you calculate the evaluation metrics?",
            "Predictions on the holdout test set were compared with true prices using MAE, RMSE, and R2.",
            "تمت مقارنة تنبؤات بيانات الاختبار مع الأسعار الحقيقية باستخدام MAE وRMSE وR2.",
            'from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score\n\nmae = mean_absolute_error(y_test, y_pred)\nrmse = mean_squared_error(y_test, y_pred, squared=False)\nr2 = r2_score(y_test, y_pred)',
        ),
        Question(
            "How do you explain feature importance?",
            "Feature importance estimates how much each input feature helped the Random Forest reduce prediction error across its trees.",
            "أهمية الخصائص تقدر مقدار مساهمة كل خاصية في مساعدة Random Forest على تقليل خطأ التنبؤ عبر الأشجار.",
            'rf_model = best_model.named_steps["model"]\nimportances = rf_model.feature_importances_\n# Top signals: freight_value, product_weight_g, description length',
        ),
        Question(
            "What result should you emphasize in the presentation?",
            "Emphasize that Random Forest reduced MAE from 57.86 BRL for the baseline to 22.54 BRL, a reduction of about 61 percent.",
            "يجب التركيز على أن Random Forest خفض MAE من 57.86 في نموذج الأساس إلى 22.54، أي انخفاض يقارب 61%.",
            'baseline_mae = 57.8563\nrf_mae = 22.5440\nreduction = (baseline_mae - rf_mae) / baseline_mae\nprint(f"{reduction:.1%}")  # about 61.0%',
        ),
    ],
}


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    add_run(
        intro,
        "Use this guide as a defense sheet. Each answer has: a short English response, an Arabic response, and code or technical evidence you can point to in the repository.",
        size=10.5,
        color=DARK,
    )

    add_key_table(
        doc,
        "Quick Project Facts",
        [
            ("Dataset", "Olist Brazilian E-Commerce Public Dataset"),
            ("Prepared modeling file", "dataset/feature_engineered_dataset.csv"),
            ("Rows used", "107,529 processed rows"),
            ("Train/test rows", "75,270 training rows / 32,259 test rows"),
            ("Target", "price, continuous item price in BRL"),
            ("Best model", "Random Forest Regressor"),
            ("Best holdout metrics", "MAE 22.54 BRL, RMSE 53.09 BRL, R2 0.775"),
            ("Saved model", "models/best_price_model.joblib"),
        ],
    )

    add_key_table(
        doc,
        "Repository Code Map",
        [
            ("Training notebook", "notebooks/03_model_training_results.ipynb"),
            ("Saved-model test notebook", "notebooks/04_test_saved_model.ipynb"),
            ("Command-line model test", "scripts/run_saved_model.py"),
            ("Model metrics", "reports/model_performance.csv and reports/model_training_metrics.json"),
            ("Cross-validation", "reports/cross_validation_results.csv"),
            ("Feature subsets", "reports/feature_subset_results.csv"),
            ("Feature importance", "reports/feature_importance.csv and reports/figures/feature_importance.png"),
            ("HTML result report", "reports/model_training_results.html"),
        ],
    )

    add_heading(doc, "General Project, Data, and Modeling Questions", level=1)
    for idx, item in enumerate(general_questions, start=1):
        add_question(doc, idx, item)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "Role-Specific Questions", level=1)
    counter = 1
    for role, questions in role_questions.items():
        add_heading(doc, role, level=2)
        for item in questions:
            add_question(doc, counter, item)
            counter += 1

    add_heading(doc, "Fast Closing Answers", level=1)
    add_label_text(
        doc,
        "If asked for the strongest sentence: ",
        "The project is a reproducible regression pipeline that uses prepared Olist transaction data to predict item price, with Random Forest achieving the best holdout result: MAE 22.54 BRL and R2 0.775.",
    )
    add_arabic_answer(
        doc,
        "أقوى إجابة مختصرة: المشروع هو مسار انحدار قابل لإعادة التشغيل يستخدم بيانات Olist المجهزة للتنبؤ بسعر المنتج، وأفضل نتيجة كانت Random Forest بقيمة MAE تساوي 22.54 وR2 تساوي 0.775."
    )
    add_code_block(
        doc,
        'Evidence to show quickly:\nreports/model_performance.csv\nmodels/best_price_model.joblib\nscripts/run_saved_model.py --rows 3',
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT_PATH)
