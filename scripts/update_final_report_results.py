from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_PATH = Path(
    r"C:\Users\gosfe\OneDrive\University\Third year\Second Semester\Machine Learning\Project Documnts\ARTI 308 Final Project Report - Retail Price Optimizer.docx"
)
METRICS_PATH = ROOT / "reports" / "model_training_metrics.json"


def fmt(value: float, digits: int = 2) -> str:
    return f"{value:,.{digits}f}"


def set_cell(cell, text: str, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(8.6)


def set_single_cell_box(table, title: str, body: str) -> None:
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)


def replace_paragraph(paragraph, text: str) -> None:
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def replace_caption(paragraph, text: str) -> None:
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(91, 99, 108)


def ensure_rows(table, total_rows: int) -> None:
    while len(table.rows) < total_rows:
        table.add_row()


def main() -> None:
    metrics = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    performance = {row["Technique"]: row for row in metrics["holdout_results"]}
    cv = {row["Technique"]: row for row in metrics["cross_validation_results"]}
    subsets = metrics["feature_subset_results"]
    best = metrics["best_model"]
    best_row = performance[best]
    best_params = metrics["best_params"]
    best_subset = min(subsets, key=lambda row: row["MAE"])

    df = pd.read_csv(ROOT / "dataset" / "feature_engineered_dataset.csv")
    stats_rows = [
        ("Price (BRL)", "price"),
        ("Freight Value (BRL)", "freight_value"),
        ("Product Weight (g)", "product_weight_g"),
        ("Description Length", "product_description_lenght"),
        ("Photos Quantity", "product_photos_qty"),
    ]
    corr_rows = [
        ("Freight Value and Price", "freight_value", "Moderate positive relationship; freight cost remains the strongest numeric price driver."),
        ("Product Weight and Price", "product_weight_g", "Moderate positive relationship; heavier products tend to be more expensive."),
        ("Product Description Length and Price", "product_description_lenght", "Weak positive relationship; richer product information has some association with price."),
        ("Product Photos Quantity and Price", "product_photos_qty", "Very weak relationship; photo count alone is not a strong predictor."),
        ("Product Volume and Price", "product_volume_cm3", "Moderate positive relationship; larger products often have higher price and logistics complexity."),
        ("Product Demand Count and Price", "product_demand_count", "Weak negative relationship; frequently ordered products are often lower-priced items."),
    ]

    doc = Document(str(REPORT_PATH))

    # Table 4 in the rendered report: repository status note.
    set_single_cell_box(
        doc.tables[4],
        "Repository status",
        (
            "The GitHub repository was cloned locally and updated from origin/data_prep. "
            "The latest feature-engineered dataset, executed model-training notebook, saved metrics, figures, model artifact, "
            "and HTML results page are available in the local project repo."
        ),
    )

    # Project deliverables: final report is now populated with actual results.
    set_cell(doc.tables[2].cell(4, 4), "Completed with training results")
    set_cell(doc.tables[3].cell(5, 0), "2240000344")


    # Table 5: statistics from the actual feature-engineered dataset.
    table = doc.tables[5]
    for i, (label, col) in enumerate(stats_rows, start=1):
        s = df[col]
        values = [label, fmt(s.mean()), fmt(s.median()), fmt(s.std()), fmt(s.max()), fmt(s.min())]
        for j, value in enumerate(values):
            set_cell(table.cell(i, j), value)

    # Table 6: correlations, expanded with engineered features.
    table = doc.tables[6]
    ensure_rows(table, len(corr_rows) + 1)
    for i, (label, col, note) in enumerate(corr_rows, start=1):
        corr = df[[col, "price"]].corr().iloc[0, 1]
        set_cell(table.cell(i, 0), label, WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(table.cell(i, 1), fmt(corr, 2))
        set_cell(table.cell(i, 2), note, WD_ALIGN_PARAGRAPH.LEFT)

    # Table 8: optimum parameters.
    table = doc.tables[8]
    param_rows = [
        ("Random Forest Regressor", "n_estimators", "80, 140", best_params["Random Forest Regressor"]["n_estimators"]),
        ("Random Forest Regressor", "max_depth", "14, None", "None" if best_params["Random Forest Regressor"]["max_depth"] is None else best_params["Random Forest Regressor"]["max_depth"]),
        ("Random Forest Regressor", "min_samples_leaf", "1", best_params["Random Forest Regressor"]["min_samples_leaf"]),
        ("Gradient Boosting Regressor", "n_estimators", "80, 140", best_params["Gradient Boosting Regressor"]["n_estimators"]),
        ("Gradient Boosting Regressor", "learning_rate", "0.05, 0.10", best_params["Gradient Boosting Regressor"]["learning_rate"]),
        ("Gradient Boosting Regressor", "max_depth", "2", best_params["Gradient Boosting Regressor"]["max_depth"]),
    ]
    for i, row in enumerate(param_rows, start=1):
        for j, value in enumerate(row):
            set_cell(table.cell(i, j), value)

    # Current result status box.
    set_single_cell_box(
        doc.tables[9],
        "Current result status",
        (
            f"Training was completed on {metrics['run_timestamp']} using the feature-engineered dataset "
            f"({metrics['dataset']['rows']:,} rows). The best holdout model was {best}, with MAE "
            f"{fmt(best_row['MAE (BRL)'])} BRL, RMSE {fmt(best_row['RMSE (BRL)'])} BRL, and R2 {best_row['R2']:.3f}."
        ),
    )

    # Table 9: holdout model results.
    table = doc.tables[10]
    ordered_models = ["Linear Regression Baseline", "Random Forest Regressor", "Gradient Boosting Regressor"]
    notes = {
        "Linear Regression Baseline": "Baseline model for comparison; weakest performance.",
        "Random Forest Regressor": "Best final model; lowest error and highest explained variance.",
        "Gradient Boosting Regressor": "Improved over baseline, but below Random Forest on this run.",
    }
    for i, name in enumerate(ordered_models, start=1):
        row = performance[name]
        values = [name, fmt(row["MAE (BRL)"]), fmt(row["RMSE (BRL)"]), f"{row['R2']:.3f}", notes[name]]
        for j, value in enumerate(values):
            set_cell(table.cell(i, j), value, WD_ALIGN_PARAGRAPH.LEFT if j in [0, 4] else WD_ALIGN_PARAGRAPH.CENTER)

    # Table 10: feature subset results.
    table = doc.tables[11]
    for i, row in enumerate(subsets, start=1):
        values = [
            row["Feature Subset"],
            row["Technique"],
            fmt(row["MAE"]),
            fmt(row["RMSE"]),
            f"{row['R2']:.3f}",
        ]
        for j, value in enumerate(values):
            set_cell(table.cell(i, j), value, WD_ALIGN_PARAGRAPH.LEFT if j in [0, 1] else WD_ALIGN_PARAGRAPH.CENTER)

    # Rubric alignment R6 now references actual results.
    set_cell(
        doc.tables[12].cell(6, 1),
        "The results section now presents measured MAE, RMSE, R2, cross-validation results, feature-selection results, and model interpretation from the executed notebook.",
        WD_ALIGN_PARAGRAPH.LEFT,
    )

    feature_selection_text = (
        "Feature selection was tested using model-based feature importance and grouped feature subsets. "
        f"The best subset in the corrected training run was {best_subset['Feature Subset']}, "
        f"which used {int(best_subset['Number of Input Columns'])} input columns and achieved MAE "
        f"{fmt(best_subset['MAE'])} BRL, RMSE {fmt(best_subset['RMSE'])} BRL, and R2 {best_subset['R2']:.3f}. "
        "This comparison shows whether the full prepared feature set or a smaller subset gives the best "
        "accuracy while keeping the feature set interpretable."
    )
    abstract_text = (
        "Digital marketplaces make pricing decisions more complex because prices depend on product attributes, "
        "freight, category demand, and marketplace behavior. This project developed a Retail Price Optimizer "
        "using the Olist Brazilian e-commerce dataset and the prepared feature-engineered CSV from the data-prep "
        "branch. The target variable is item price in BRL. The corrected model-training branch tested Linear "
        "Regression, Random Forest Regressor, and Gradient Boosting Regressor after excluding target-derived "
        f"leakage columns. {best} produced the best holdout performance with MAE "
        f"{fmt(best_row['MAE (BRL)'])} BRL, RMSE {fmt(best_row['RMSE (BRL)'])} BRL, and R2 "
        f"{best_row['R2']:.3f}. The results show that logistics, product scale, demand-count, and category "
        "signals can support more evidence-based pricing decisions."
    )
    first_experiment_text = (
        f"The first experiment used the complete cleaned and feature-engineered set after excluding "
        f"target-derived leakage columns. {best} clearly outperformed the Linear Regression baseline: "
        f"MAE decreased from {fmt(performance['Linear Regression Baseline']['MAE (BRL)'])} BRL to "
        f"{fmt(best_row['MAE (BRL)'])} BRL, RMSE decreased from "
        f"{fmt(performance['Linear Regression Baseline']['RMSE (BRL)'])} BRL to "
        f"{fmt(best_row['RMSE (BRL)'])} BRL, and R2 increased from "
        f"{performance['Linear Regression Baseline']['R2']:.3f} to {best_row['R2']:.3f}."
    )
    conclusion_text = (
        f"This report presents a Retail Price Optimizer for Brazilian e-commerce using machine learning "
        f"regression. The project addresses the limitations of static pricing by implementing a data-driven "
        f"pipeline that merges marketplace transaction data, uses the prepared feature-engineered dataset, "
        f"and compares Linear Regression, Random Forest, and Gradient Boosting models. The executed results "
        f"selected {best} as the final model, achieving MAE {fmt(best_row['MAE (BRL)'])} BRL, RMSE "
        f"{fmt(best_row['RMSE (BRL)'])} BRL, and R2 {best_row['R2']:.3f} on the holdout test set. "
        "The notebook and HTML report provide the metrics, figures, feature importance, and model artifact "
        "needed to reproduce the result."
    )

    replacements = {
        "Machine learning expands this research direction by allowing retailers to model complex relationships that are difficult to capture using simple linear assumptions. Random Forest and Gradient Boosting are especially suitable for structured retail data because they learn nonlinear feature interactions without requiring the analyst to manually specify every interaction. For this project, the Olist dataset provides a realistic e-commerce context where price is influenced by freight, product weight, product category, product information, order timing, and seller-related factors.": (
            "Machine learning expands this research direction by allowing retailers to model complex relationships that are difficult to capture using simple linear assumptions. Random Forest and Gradient Boosting are especially suitable for structured retail data because they learn nonlinear feature interactions without requiring the analyst to manually specify every interaction. For this project, the Olist dataset provides a realistic e-commerce context where price is influenced by freight value, product weight, product category, product information, product volume, and demand-count features."
        ),
        "The dataset used in this project is the Brazilian E-Commerce Public Dataset by Olist. It contains approximately 100,000 anonymized orders from Brazilian marketplaces between 2016 and 2018. The project uses relational files such as order items, orders, products, sellers, customers, and product category translation. The target attribute is item price in BRL. Candidate predictive attributes include freight value, product weight, product dimensions, product category, product description length, number of photos, order timing, seller/customer location features, and engineered logistics features.": (
            "The dataset used in this project is the Brazilian E-Commerce Public Dataset by Olist. It contains approximately 100,000 anonymized orders from Brazilian marketplaces between 2016 and 2018. The model-training notebook uses the prepared feature-engineered CSV from the data-prep branch. The target attribute is item price in BRL. Candidate predictive attributes include freight value, product weight, product dimensions, product category, product description length, number of photos, product volume, product demand count, category demand count, freight level, and product size level."
        ),
        "The expected contribution is a reproducible machine learning workflow that helps retailers move from guess-based pricing to evidence-based decisions that balance market acceptance and revenue potential.": (
            f"The executed model-training run shows that {best} achieved the strongest performance, with MAE "
            f"{fmt(best_row['MAE (BRL)'])} BRL, RMSE {fmt(best_row['RMSE (BRL)'])} BRL, and R2 {best_row['R2']:.3f}, "
            "supporting evidence-based pricing decisions for the prepared Olist dataset."
        ),
        "Hyperparameter optimization will be performed using GridSearchCV with cross-validation. The search objective is to minimize validation MAE while also checking RMSE and R2 so the selected model is accurate and stable. Random state values should be fixed to make results reproducible.": (
            "Hyperparameter optimization was performed using GridSearchCV. The search minimized validation MAE while also checking RMSE and R2 for stability. Random state values were fixed for reproducibility. The selected Random Forest configuration used 140 trees with no maximum depth and min_samples_leaf of 1, while Gradient Boosting used 140 estimators, learning_rate of 0.10, and max_depth of 2."
        ),
        "The first experiment will use the complete feature set after cleaning and encoding. The purpose of this experiment is to compare whether ensemble models outperform the Linear Regression baseline. A successful result would show lower MAE and RMSE for Random Forest or Gradient Boosting, along with a higher R2 value. Feature-importance analysis will then be used to identify the most influential variables, with freight value, product weight, product category, and product dimensions expected to be important predictors based on the statistical analysis.": (
            first_experiment_text
        ),
        "Feature selection is included to test whether a smaller and more interpretable subset of features can maintain or improve performance. The initial feature ranking will use correlation analysis for numeric variables and model-based feature importance from the tree ensembles. Recursive feature elimination can then be applied by reducing the feature set and comparing cross-validation results.": (
            feature_selection_text
        ),
        "Feature selection was tested using model-based feature importance and grouped feature subsets. The best subset was the top numeric predictors plus product category, which used only 9 input columns and improved the Random Forest holdout result to MAE 22.87 BRL, RMSE 54.31 BRL, and R2 0.765. This indicates that a smaller, better-focused feature set can be more accurate and easier to interpret than using every engineered column.": (
            feature_selection_text
        ),
        "After training, the final model should be selected using a balanced interpretation of MAE, RMSE, and R2. MAE gives the clearest business interpretation because it shows the expected average pricing error in BRL. RMSE is used to check whether the model makes large mistakes on expensive or unusual products. R2 shows whether the selected predictors explain a meaningful share of price variation. The final discussion should include a predicted-versus-actual plot, residual analysis, and feature-importance plot.": (
            f"The final model was selected using MAE, RMSE, and R2. {best} was selected because it produced the lowest holdout error and highest explained variance among the tested models. The 10-fold validation sample also favored Random Forest, with mean CV MAE {fmt(cv['Random Forest Regressor']['CV MAE Mean'])} BRL compared with {fmt(cv['Gradient Boosting Regressor']['CV MAE Mean'])} BRL for Gradient Boosting. The notebook includes predicted-versus-actual, residual, and feature-importance plots."
        ),
        "The expected business interpretation is that the optimizer will help retailers choose prices based on historical evidence rather than guessing. If the model identifies freight value and product weight as important, the retailer can better understand how logistics costs influence the recommended price. If product category and description features are important, the retailer can adjust pricing strategy by category rather than using one fixed markup policy for all products.": (
            "The feature-importance analysis supports the business interpretation. The strongest predictors were freight value, product weight, product description length, product demand count, and product dimensions. This means the optimizer is learning from logistics cost, product scale, demand frequency, and category-related signals, which can help retailers avoid one fixed markup policy and make more data-driven category-level pricing decisions."
        ),
        "This report presents a Retail Price Optimizer for Brazilian e-commerce using machine learning regression. The project addresses the limitations of static pricing by proposing a data-driven pipeline that merges marketplace transaction data, engineers product and logistics features, and compares Random Forest and Gradient Boosting models. The selected performance measures and optimization strategy make the experiment reproducible and interpretable. The final training step will complete the metric tables, determine the best model, and produce the final feature-importance and predicted-versus-actual plots.": (
            conclusion_text
        ),
    }
    prefix_replacements = [
        ("Digital marketplaces make pricing decisions more complex", abstract_text),
        ("Digital marketplaces make pricing decisions more difficult", abstract_text),
        ("The first experiment used the complete cleaned", first_experiment_text),
        ("This report presents a Retail Price Optimizer", conclusion_text),
    ]

    for paragraph in doc.paragraphs:
        text = paragraph.text
        for old, new in replacements.items():
            if old in text:
                replace_paragraph(paragraph, text.replace(old, new))
                text = paragraph.text
        for prefix, new in prefix_replacements:
            if text.startswith(prefix) and text != new:
                replace_paragraph(paragraph, new)
                text = paragraph.text
        if "Table 8. Hyperparameter search plan for proposed models" in text:
            replace_caption(paragraph, "Table 8. Hyperparameter search results for proposed models")
        elif "Table 10. Planned comparison of feature subsets" in text:
            replace_caption(paragraph, "Table 10. Feature subset comparison results")
        elif "Table 8. Hyperparameter search results for proposed models" in text:
            replace_caption(paragraph, "Table 8. Hyperparameter search results for proposed models")
        elif "Table 10. Feature subset comparison results" in text:
            replace_caption(paragraph, "Table 10. Feature subset comparison results")

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(REPORT_PATH))
    print(f"Updated {REPORT_PATH}")


if __name__ == "__main__":
    main()
